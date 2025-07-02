import tempfile
from io import BytesIO
from zipfile import BadZipFile

from core.enums.trasnaltes import FileExtensionType, TranslationType
from docx import Document
from fastapi import HTTPException, UploadFile
from fastapi.responses import FileResponse
from starlette import status

from services.translation_api import TranslationAPI


def _check_file_ext(filename: str, ext: FileExtensionType) -> bool:
    if not filename:
        return False

    parsed_ext = filename.split(".")[-1].lower()
    return parsed_ext == ext.value.lower()


class TranslatesService:
    def __init__(self, translation_api: TranslationAPI) -> None:
        self._translation_api = translation_api

    async def translate_docx(self, file: UploadFile, translation_type: TranslationType) -> FileResponse:
        if not _check_file_ext(filename=file.filename, ext=FileExtensionType.DOCX):
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid file extension")

        file_stream = BytesIO(await file.read())

        try:
            doc = Document(file_stream)
        except BadZipFile:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid document")

        doc = await self._translate_paragraphs(doc, translation_type)
        doc = await self._translate_tables(doc, translation_type)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            doc.save(temp_file.name)
            temp_path = temp_file.name

        return FileResponse(
            path=temp_path,
            filename=f"translated_{file.filename}",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    async def _translate_paragraphs(self, document: Document, translation_type: TranslationType) -> Document:
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                for run in paragraph.runs:
                    if run.text.strip():
                        translated = await self._translation_api.translate(
                            translation_type=translation_type, text=run.text
                        )
                        run.text = translated.translated_text

        return document

    async def _translate_tables(self, document: Document, translation_type: TranslationType) -> Document:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip():
                                translated = await self._translation_api.translate(
                                    translation_type=translation_type, text=run.text
                                )
                                run.text = translated.translated_text

        return document
